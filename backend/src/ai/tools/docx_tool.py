"""
DOCX Tool — Read, Create, and Update Word documents for AI Agents.

Uses python-docx to manipulate .docx files. Saves generated files
via ArtifactService for traceability.
"""

import json
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

from src.ai.tools.base import Tool

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None
    Inches = None
    Pt = None
    WD_ALIGN_PARAGRAPH = None

# Artifact root for system-generated files
BASE_DIR = Path(__file__).resolve().parents[3] / "artifact" / "system-generated"


class DocxTool(Tool):
    """Tool for reading, creating, and updating Word (.docx) documents.

    Actions:
        create  — Build a new docx from structured content (title, sections, tables)
        read    — Extract text, tables, and metadata from an existing docx
        update  — Open an existing docx and append/replace content
    """

    name = "docx_tool"
    description = (
        "Read, create, or update Word (.docx) documents. "
        "Input: JSON with 'action' ('create', 'read', or 'update') "
        "and action-specific parameters. "
        "For 'create': provide 'filename', 'title', and 'sections' array. "
        "For 'read': provide 'file_path'. "
        "For 'update': provide 'file_path' and content to append/replace."
    )

    async def run(self, input_data: str) -> str:
        return await self.run_with_context(input_data, context=None)

    async def run_with_context(self, input_data: str, context=None) -> str:
        # Lazy re-import: picks up python-docx if installed after worker started
        global Document, Inches, Pt, WD_ALIGN_PARAGRAPH
        if Document is None:
            try:
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                raise RuntimeError(
                    "python-docx is not installed. Install via: pip install python-docx"
                )

        try:
            params = json.loads(input_data) if isinstance(input_data, str) else input_data
            action = params.get("action")

            if action == "create":
                return await self._create(params, context)
            elif action == "read":
                return self._read(params)
            elif action == "update":
                return await self._update(params, context)
            else:
                return json.dumps({"error": f"Unknown action: {action}. Use 'create', 'read', or 'update'."})
        except Exception as e:
            return json.dumps({"error": f"DocxTool error: {str(e)}"})

    # ------------------------------------------------------------------
    # CREATE
    # ------------------------------------------------------------------

    async def _create(self, params: Dict[str, Any], context: Optional[Dict[str, Any]]) -> str:
        doc = Document()

        title = params.get("title", "Untitled Document")
        filename = params.get("filename", "document.docx")
        if not filename.endswith(".docx"):
            filename += ".docx"

        # Title
        doc.add_heading(title, level=0)

        # Sections
        sections = params.get("sections", [])
        for section in sections:
            heading = section.get("heading")
            if heading:
                level = section.get("level", 1)
                doc.add_heading(heading, level=min(level, 9))

            # Paragraphs
            for para_text in section.get("paragraphs", []):
                doc.add_paragraph(para_text)

            # Bullet list
            for bullet in section.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

            # Numbered list
            for item in section.get("numbered_items", []):
                doc.add_paragraph(item, style="List Number")

            # Table
            table_data = section.get("table")
            if table_data:
                headers = table_data.get("headers", [])
                rows = table_data.get("rows", [])
                if headers:
                    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
                    tbl.style = "Table Grid"
                    for i, h in enumerate(headers):
                        tbl.rows[0].cells[i].text = str(h)
                    for r_idx, row in enumerate(rows):
                        for c_idx, cell in enumerate(row):
                            if c_idx < len(headers):
                                tbl.rows[r_idx + 1].cells[c_idx].text = str(cell)

        # Save
        company_id = (context or {}).get("company_id", params.get("company_id", "default"))
        date_str = datetime.utcnow().strftime("%Y-%m-%d")
        output_dir = BASE_DIR / str(company_id) / date_str
        output_dir.mkdir(parents=True, exist_ok=True)

        file_path = output_dir / filename
        doc.save(str(file_path))

        # Register artifact
        artifact_id = await self._register_artifact(file_path, filename, company_id, params)

        result = {
            "status": "success",
            "file_path": str(file_path),
            "message": f"Created {filename} with {len(sections)} sections",
        }
        if artifact_id:
            result["artifact_id"] = artifact_id
            result["download_url"] = f"/api/v1/artifacts/{artifact_id}/download"
            result["document_path"] = f"/api/v1/artifacts/{artifact_id}/download"

        return json.dumps(result)

    # ------------------------------------------------------------------
    # READ
    # ------------------------------------------------------------------

    def _read(self, params: Dict[str, Any]) -> str:
        file_path = params.get("file_path", "")
        if not file_path:
            return json.dumps({"error": "Missing 'file_path'"})

        if not os.path.exists(file_path):
            fallback = os.path.join("uploads", os.path.basename(file_path))
            if os.path.exists(fallback):
                file_path = fallback
            else:
                return json.dumps({"error": f"File not found: {file_path}"})

        doc = Document(file_path)

        # Extract metadata
        core = doc.core_properties
        metadata = {
            "title": core.title or "",
            "author": core.author or "",
            "created": str(core.created) if core.created else "",
            "modified": str(core.modified) if core.modified else "",
        }

        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append({
                "text": para.text,
                "style": para.style.name if para.style else "",
            })

        # Extract tables
        tables = []
        for tbl_idx, tbl in enumerate(doc.tables):
            rows = []
            for row in tbl.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append({"table_index": tbl_idx, "rows": rows})

        # Limit output for LLM context
        limit = params.get("limit", 200)
        if len(paragraphs) > limit:
            paragraphs = paragraphs[:limit]

        return json.dumps({
            "status": "success",
            "metadata": metadata,
            "paragraphs": paragraphs,
            "tables": tables,
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables),
        })

    # ------------------------------------------------------------------
    # UPDATE
    # ------------------------------------------------------------------

    async def _update(self, params: Dict[str, Any], context: Optional[Dict[str, Any]]) -> str:
        file_path = params.get("file_path", "")
        if not file_path or not os.path.exists(file_path):
            return json.dumps({"error": f"File not found: {file_path}"})

        doc = Document(file_path)

        changes_made = 0

        # Append heading
        heading = params.get("append_heading")
        if heading:
            level = params.get("heading_level", 1)
            doc.add_heading(heading, level=min(level, 9))
            changes_made += 1

        # Append paragraphs
        for para_text in params.get("append_paragraphs", []):
            doc.add_paragraph(para_text)
            changes_made += 1

        # Append bullets
        for bullet in params.get("append_bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")
            changes_made += 1

        # Append table
        table_data = params.get("append_table")
        if table_data:
            headers = table_data.get("headers", [])
            rows = table_data.get("rows", [])
            if headers:
                tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
                tbl.style = "Table Grid"
                for i, h in enumerate(headers):
                    tbl.rows[0].cells[i].text = str(h)
                for r_idx, row in enumerate(rows):
                    for c_idx, cell in enumerate(row):
                        if c_idx < len(headers):
                            tbl.rows[r_idx + 1].cells[c_idx].text = str(cell)
                changes_made += 1

        # Find and replace text
        replacements = params.get("replacements", [])
        for repl in replacements:
            old_text = repl.get("find", "")
            new_text = repl.get("replace", "")
            if old_text:
                for para in doc.paragraphs:
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                changes_made += 1

        # Save (overwrite or save-as)
        save_path = params.get("save_as", file_path)
        doc.save(save_path)

        return json.dumps({
            "status": "success",
            "file_path": save_path,
            "changes_made": changes_made,
            "message": f"Updated document with {changes_made} changes",
        })

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    async def _register_artifact(
        file_path: Path, filename: str, company_id: str, params: Dict[str, Any]
    ) -> str | None:
        """Register file as an artifact. Returns artifact ID or None."""
        if company_id and company_id != "default":
            try:
                from uuid import UUID as _UUID
                from src.common.database import AsyncSessionLocal
                from src.ai.artifact_service import ArtifactService, ORIGIN_SYSTEM

                async with AsyncSessionLocal() as _db:
                    art_svc = ArtifactService(_db)
                    with open(file_path, "rb") as f:
                        file_bytes = f.read()
                    artifact = await art_svc.save_artifact(
                        file_bytes=file_bytes,
                        file_name=filename,
                        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        file_category="documents",
                        origin=ORIGIN_SYSTEM,
                        company_id=_UUID(str(company_id)),
                        purpose=params.get("purpose", "AI-generated Word document"),
                        generated_by=params.get("generated_by", "docx_tool"),
                    )
                    return str(artifact.id)
            except Exception:
                pass  # Non-fatal
        return None

    def get_function_schema(self) -> Dict[str, Any]:
        return {
            "name": self.name,
            "description": self.description,
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["create", "read", "update"],
                        "description": "Action to perform",
                    },
                    "filename": {
                        "type": "string",
                        "description": "Output filename for 'create' (e.g. report.docx)",
                    },
                    "title": {
                        "type": "string",
                        "description": "Document title for 'create'",
                    },
                    "sections": {
                        "type": "array",
                        "description": (
                            "Array of section objects for 'create'. Each section can have: "
                            "'heading' (str), 'level' (int 1-9), 'paragraphs' (str[]), "
                            "'bullets' (str[]), 'numbered_items' (str[]), "
                            "'table' ({headers: str[], rows: str[][]})"
                        ),
                        "items": {"type": "object"},
                    },
                    "file_path": {
                        "type": "string",
                        "description": "Path to existing docx for 'read' or 'update'",
                    },
                    "append_heading": {
                        "type": "string",
                        "description": "Heading to append (for 'update')",
                    },
                    "heading_level": {
                        "type": "integer",
                        "description": "Heading level 1-9 (for 'update', default: 1)",
                    },
                    "append_paragraphs": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Paragraphs to append (for 'update')",
                    },
                    "append_bullets": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Bullet items to append (for 'update')",
                    },
                    "append_table": {
                        "type": "object",
                        "description": "Table to append: {headers: str[], rows: str[][]} (for 'update')",
                    },
                    "replacements": {
                        "type": "array",
                        "description": "Find-and-replace pairs: [{find: str, replace: str}] (for 'update')",
                        "items": {"type": "object"},
                    },
                    "save_as": {
                        "type": "string",
                        "description": "Optional save-as path (for 'update')",
                    },
                    "limit": {
                        "type": "integer",
                        "description": "Max paragraphs to return (for 'read', default: 200)",
                    },
                },
                "required": ["action"],
            },
        }
